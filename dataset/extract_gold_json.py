"""
Extracts gold-standard ExperimentReport JSON from existing, already-correct
example .docx files (E01-E03 style). This is NOT a general-purpose docx
parser — it's a heuristic extractor tuned to the specific pattern observed
across the example set:

    - A title paragraph ("EXPERIMENT NN" or "EXPERIMENT NO. N")
    - Particulars fields, either as prose ("Aim: ...", "Author: ... (Section:
      A2, Roll No: 33)") or as a table (label | value rows)
    - Numbered body section headings ("2. Short Description...", bold, short)
    - Bullet items starting with "\u2022 ", optionally with a "Lead-in: rest"
      pattern
    - Inline images followed by a "Fig. N \u2014 caption" paragraph

Output is renumbered canonically: Particulars is always implicit Section 1
(never numbered in the source), and body sections are renumbered 2, 3, 4...
in document order regardless of what number the source document used. This
is a deliberate normalization, not a parsing artifact — the whole point of
the pipeline is a STANDARDIZED output structure; source documents were not
all self-consistent about numbering (E02 numbers particulars as "1." and
starts body at "2."; E03 leaves particulars unnumbered and starts body at
"1."). The schema's own validator (schema.py) enforces the canonical
2,3,4... convention, so this extractor must match it.

Every extracted result is run through ExperimentReport.model_validate()
before being written — if extraction produced something that violates the
schema, this script errors instead of writing bad gold data.

KNOWN LIMITATION: this heuristic set does NOT capture E01a's structure
(multi-level numbered subsections like "2.1", ALL-CAPS unnumbered headings
like "SUMMARY OF REQUIREMENTS", deeply nested bullet hierarchies with
bold requirement codes like "FR-1.1"). E01a does not fit the current
ExperimentReport schema and is excluded from this extraction pass pending
a decision on whether to extend the schema or treat it as an exception
(same as E01b/the embedded SRS).

Usage:
    python extract_gold_json.py ../examples/E02_5A2_33.docx --experiment-no 2 \
        --semester 5 --division A2 --roll 33 --out target_json/E02.json
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from datetime import datetime

import docx
from docx.oxml.ns import qn

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "schema"))
from schema import ExperimentReport  # noqa: E402


SECTION_HEADING_RE = re.compile(r"^\s*(\d+)\.\s+(.+?)\s*$")
BULLET_PREFIX_RE = re.compile(r"^[\u2022\-]\s*")
LEAD_IN_RE = re.compile(r"^([^:]{3,60}):\s+(.+)$")
NUMBERED_LEAD_IN_RE = re.compile(r"^\d+(?:\.\d+)+\s+([^:]{3,80}):\s+(.+)$")

PARTICULARS_LABELS = {
    "case_study_title": re.compile(r"^case study title\s*:\s*(.+)$", re.I),
    "aim": re.compile(r"^aim(?:\s+of\s+experiment)?\s*:\s*(.+)$", re.I),
    "problem_statement": re.compile(r"^problem statement\s*:\s*(.+)$", re.I),
    "author": re.compile(r"^author\s*:\s*(.+)$", re.I),
    "date_of_compilation": re.compile(
        r"^date\s+of\s+(?:compilation|experiment)(?:/submission)?\s*:\s*(.+)$", re.I
    ),
}
INLINE_SECTION_ROLL_RE = re.compile(
    r"\(?\s*Section\s*:\s*([A-Za-z0-9]+)\s*,\s*Roll\s*(?:No\.?|Number)\s*:\s*(\d+)\s*\)?", re.I
)


def _paragraph_has_image(paragraph) -> bool:
    return len(paragraph._p.findall(".//" + qn("w:drawing"))) > 0


def _extract_image(paragraph, doc, out_dir: str, idx: int) -> str | None:
    drawings = paragraph._p.findall(".//" + qn("w:drawing"))
    if not drawings:
        return None
    blips = drawings[0].findall(".//" + qn("a:blip"))
    if not blips:
        return None
    rId = blips[0].get(qn("r:embed"))
    if rId is None or rId not in doc.part.rels:
        return None
    image_part = doc.part.rels[rId].target_part
    ext = image_part.content_type.split("/")[-1]
    os.makedirs(out_dir, exist_ok=True)
    filename = f"figure_{idx}.{ext}"
    path = os.path.join(out_dir, filename)
    with open(path, "wb") as f:
        f.write(image_part.blob)
    return path


def _try_parse_date(text: str) -> str:
    text = text.strip().rstrip(".")
    for fmt in ("%B %d, %Y", "%d %B %Y", "%Y-%m-%d", "%d/%m/%Y"):
        try:
            return datetime.strptime(text, fmt).date().isoformat()
        except ValueError:
            continue
    return text  # let pydantic raise if still unparseable


def extract(docx_path: str, experiment_no: int, semester: str, division: str,
            roll: str, image_out_dir: str, part_suffix: str | None = None) -> dict:
    doc = docx.Document(docx_path)

    particulars_raw = {
        "case_study_title": None,
        "aim": None,
        "problem_statement": None,
        "author": None,
        "section": division,
        "roll_number": roll,
        "date_of_compilation": None,
    }

    sections = []
    current_section = None
    pending_bullets = None
    fig_counter = 0
    title_text = f"Experiment {experiment_no}"
    next_section_number = 2  # canonical renumbering, per module docstring

    def flush_bullets():
        nonlocal pending_bullets
        if pending_bullets and current_section is not None:
            current_section["content"].append(
                {"type": "bullet_list", "items": pending_bullets}
            )
        pending_bullets = None

    # --- particulars table (E02-style) ---
    table_particulars = {}
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = row.cells[0].text.strip().lower()
            value = row.cells[1].text.strip()
            if "case study" in label:
                table_particulars["case_study_title"] = value
            elif "aim" in label:
                table_particulars["aim"] = value
            elif "problem statement" in label:
                table_particulars["problem_statement"] = value
            elif label == "author":
                table_particulars["author"] = value
            elif "roll" in label:
                table_particulars["roll_number"] = value
            elif label == "section" or label == "division":
                table_particulars["section"] = value
            elif "date" in label:
                table_particulars["date_of_compilation"] = value
    particulars_raw.update({k: v for k, v in table_particulars.items() if v})

    body_started = False

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()

        if _paragraph_has_image(p):
            fig_counter += 1
            img_path = _extract_image(p, doc, image_out_dir, fig_counter)
            caption = ""
            # look ahead for a caption paragraph starting with "Fig."
            if i + 1 < len(doc.paragraphs):
                nxt = doc.paragraphs[i + 1].text.strip()
                if nxt.lower().startswith("fig"):
                    caption = nxt
            # look back for a manual "Figure: <description>" lead-in line
            # (seen in E03) and merge it into the caption if the forward
            # caption was label-only (e.g. bare "Fig. 1")
            if i - 1 >= 0:
                prev = doc.paragraphs[i - 1].text.strip()
                fig_lead = re.match(r"^figure\s*:\s*(.+)$", prev, re.I)
                if fig_lead and (not caption or len(caption) < 15):
                    label = caption or f"Fig. {fig_counter}"
                    caption = f"{label} \u2014 {fig_lead.group(1).strip()}"
            if current_section is not None and img_path:
                flush_bullets()
                current_section["content"].append(
                    {"type": "figure", "image_ref": img_path,
                     "caption": caption or f"Fig. {fig_counter}",
                     "figure_number": fig_counter}
                )
            continue

        if not text:
            continue

        # skip a caption line already consumed above
        if text.lower().startswith("fig.") or text.lower().startswith("fig "):
            continue
        # skip a manual "Figure: <caption>" lead-in line that sits before
        # the actual inline image (seen in E03) — redundant with the figure
        # block's own caption, not useful as a standalone paragraph
        if re.match(r"^figure\s*:\s*.+$", text, re.I) and i + 1 < len(doc.paragraphs):
            continue

        heading_match = SECTION_HEADING_RE.match(text)
        is_heading = bool(heading_match) and len(text) < 100 and (
            p.runs and p.runs[0].bold
        )

        if not body_started:
            # still in the pre-body particulars zone (unless this line IS
            # already a heading, e.g. E03 has no "1. Particulars" heading
            # and jumps straight to "1. Short Description...")
            matched_field = False
            for field, rx in PARTICULARS_LABELS.items():
                m = rx.match(text)
                if m:
                    value = m.group(1).strip()
                    inline = INLINE_SECTION_ROLL_RE.search(value)
                    if inline:
                        particulars_raw["section"] = inline.group(1)
                        particulars_raw["roll_number"] = inline.group(2)
                        value = INLINE_SECTION_ROLL_RE.sub("", value).strip()
                    particulars_raw[field] = value
                    matched_field = True
                    break
            if matched_field:
                continue
            if text.upper().startswith("EXPERIMENT"):
                title_text = text
                continue
            if re.match(r"^1\.\s+particulars", text, re.I):
                continue  # the "1. Particulars" heading itself, not content
            if not is_heading:
                # stray prose before body starts (e.g. E01a's repeated title
                # line) — ignore rather than risk mis-filing it
                continue

        # From here on we're in the body.
        if is_heading:
            body_started = True
            flush_bullets()
            title_only = heading_match.group(2)
            current_section = {
                "number": next_section_number,
                "title": title_only,
                "content": [],
            }
            next_section_number += 1
            sections.append(current_section)
            continue

        if current_section is None:
            continue  # nothing to attach this text to yet

        bullet_match = BULLET_PREFIX_RE.match(text)
        numbered_lead_match = NUMBERED_LEAD_IN_RE.match(text)
        if bullet_match or numbered_lead_match:
            body_started = True
            if pending_bullets is None:
                pending_bullets = []
            if bullet_match:
                item_text = BULLET_PREFIX_RE.sub("", text)
                lead_match = LEAD_IN_RE.match(item_text)
                if lead_match:
                    pending_bullets.append(
                        {"lead_in": lead_match.group(1), "text": lead_match.group(2)}
                    )
                else:
                    pending_bullets.append({"lead_in": None, "text": item_text})
            else:
                pending_bullets.append(
                    {"lead_in": numbered_lead_match.group(1), "text": numbered_lead_match.group(2)}
                )
            continue

        flush_bullets()
        current_section["content"].append({"type": "paragraph", "text": text})

    flush_bullets()

    for k in ("case_study_title", "aim", "problem_statement", "author", "date_of_compilation"):
        if not particulars_raw[k]:
            raise ValueError(f"Could not extract particulars field '{k}' from {docx_path}. "
                              f"Manual fixup needed.")

    particulars_raw["date_of_compilation"] = _try_parse_date(particulars_raw["date_of_compilation"])

    instance = {
        "title": title_text,
        "particulars": particulars_raw,
        "sections": sections,
        "submission_meta": {
            "experiment_number": experiment_no,
            "semester_prefix": semester,
            "division": division,
            "roll_number": roll,
            "part_suffix": part_suffix,
        },
    }
    return instance


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx_path")
    ap.add_argument("--experiment-no", type=int, required=True)
    ap.add_argument("--semester", default="5")
    ap.add_argument("--division", required=True)
    ap.add_argument("--roll", required=True)
    ap.add_argument("--part-suffix", default=None)
    ap.add_argument("--out", required=True)
    ap.add_argument("--image-dir", default=None,
                     help="Defaults to <out_dir>/assets/")
    args = ap.parse_args()

    out_dir = os.path.dirname(args.out) or "."
    image_dir = args.image_dir or os.path.join(out_dir, "assets")

    raw_instance = extract(
        args.docx_path, args.experiment_no, args.semester, args.division,
        args.roll, image_dir, args.part_suffix,
    )

    # Hard validation gate: refuse to write anything that doesn't conform.
    report = ExperimentReport.model_validate(raw_instance)

    os.makedirs(out_dir, exist_ok=True)
    with open(args.out, "w") as f:
        json.dump(report.model_dump(mode="json"), f, indent=2)

    print(f"OK: wrote validated gold JSON -> {args.out}")
    print(f"  Sections extracted: {[s.heading for s in report.sections]}")


if __name__ == "__main__":
    main()
