"""
Deterministic formatter: ExperimentReport (validated pydantic instance) -> DOCX.

This module has ZERO content-generation logic. Every parameter here is a
direct, literal application of Compilation_Guidelines.pdf as mapped in
formatting_spec.md. If the visual output is wrong, the fix belongs in this
file — never in the LLM prompt.

Usage:
    from schema import ExperimentReport
    from formatter import render_docx

    report = ExperimentReport.model_validate(json_from_llm)
    render_docx(report, "output.docx")
"""

from __future__ import annotations

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
import os

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "schema"))

from schema import (
    ExperimentReport,
    Particulars,
    Section,
    ParagraphBlock,
    BulletListBlock,
    FigureBlock,
    TableBlock,
)

# ---------------------------------------------------------------------------
# Constants pulled straight from the guideline (formatting_spec.md)
# ---------------------------------------------------------------------------

FONT_NAME = "Times New Roman"
TITLE_SIZE = Pt(14)
SECTION_HEADING_SIZE = Pt(12)
NORMAL_SIZE = Pt(12)

LINE_SPACING = 1.25
SPACE_BEFORE = Pt(0)
SPACE_AFTER = Pt(6)

MARGIN_LEFT = Mm(37)
MARGIN_RIGHT = Mm(25)
MARGIN_TOP = Mm(25)
MARGIN_BOTTOM = Mm(25)

PAGE_WIDTH = Mm(210)
PAGE_HEIGHT = Mm(297)

# Usable content width for images = page width - left margin - right margin
CONTENT_WIDTH = Mm(210 - 37 - 25)  # 148 mm


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _set_font(run, size, bold=False, italic=False):
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    # Ensure the east-asian font slot doesn't silently fall back to a
    # different face for any non-ASCII characters that slip through.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), FONT_NAME)


def _apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = paragraph.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = SPACE_BEFORE
    pf.space_after = SPACE_AFTER
    pf.alignment = alignment


def _add_normal_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    _apply_paragraph_format(p, alignment)
    run = p.add_run(text)
    _set_font(run, NORMAL_SIZE, bold=False)
    return p


def _add_labelled_paragraph(doc, label, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Renders 'Label: text' with only the label bolded — used for both
    Particulars fields and bulleted lead-ins, matching the pattern observed
    across E01–E03."""
    p = doc.add_paragraph()
    _apply_paragraph_format(p, alignment)
    label_run = p.add_run(f"{label}: ")
    _set_font(label_run, NORMAL_SIZE, bold=True)
    text_run = p.add_run(text)
    _set_font(text_run, NORMAL_SIZE, bold=False)
    return p


def _add_page_number_footer(section):
    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    _set_font(run, NORMAL_SIZE, bold=False)

    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


# ---------------------------------------------------------------------------
# Block renderers
# ---------------------------------------------------------------------------

def _render_paragraph_block(doc, block: ParagraphBlock):
    _add_normal_paragraph(doc, block.text)


def _render_bullet_list_block(doc, block: BulletListBlock):
    for item in block.items:
        p = doc.add_paragraph(style="List Bullet")
        _apply_paragraph_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
        if item.lead_in:
            lead_run = p.add_run(f"{item.lead_in}: ")
            _set_font(lead_run, NORMAL_SIZE, bold=True)
        text_run = p.add_run(item.text)
        _set_font(text_run, NORMAL_SIZE, bold=False)


def _render_figure_block(doc, block: FigureBlock):
    try:
        doc.add_picture(block.image_ref, width=CONTENT_WIDTH)
        pic_para = doc.paragraphs[-1]
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        # Missing/unreadable image asset must not silently corrupt the
        # document — surface a visible placeholder instead of failing the
        # whole render, since this runs unattended in a pipeline.
        placeholder = doc.add_paragraph()
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = placeholder.add_run(f"[MISSING IMAGE: {block.image_ref} ({e})]")
        _set_font(run, NORMAL_SIZE, italic=True)

    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_paragraph_format(caption_para, WD_ALIGN_PARAGRAPH.CENTER)
    caption_run = caption_para.add_run(block.caption)
    _set_font(caption_run, NORMAL_SIZE, italic=True)


def _render_table_block(doc, block: TableBlock):
    n_cols = len(block.rows[0])
    n_rows = len(block.rows) + (1 if block.header else 0)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row_offset = 0
    if block.header:
        for c, val in enumerate(block.header):
            cell = table.rows[0].cells[c]
            run = cell.paragraphs[0].add_run(val)
            _set_font(run, NORMAL_SIZE, bold=True)
        row_offset = 1

    for r, row_vals in enumerate(block.rows):
        for c, val in enumerate(row_vals):
            cell = table.rows[r + row_offset].cells[c]
            run = cell.paragraphs[0].add_run(val)
            _set_font(run, NORMAL_SIZE, bold=False)


_BLOCK_RENDERERS = {
    ParagraphBlock: _render_paragraph_block,
    BulletListBlock: _render_bullet_list_block,
    FigureBlock: _render_figure_block,
    TableBlock: _render_table_block,
}


def _render_block(doc, block):
    renderer = _BLOCK_RENDERERS.get(type(block))
    if renderer is None:
        raise ValueError(f"No renderer registered for block type {type(block)}")
    renderer(doc, block)


# ---------------------------------------------------------------------------
# Section-level renderers
# ---------------------------------------------------------------------------

def _render_particulars(doc, particulars: Particulars):
    """Particulars is ALWAYS rendered as prose (decided over table layout —
    see formatting_spec.md), as Section 1, with each field as a bolded-label
    paragraph."""
    heading_p = doc.add_paragraph()
    _apply_paragraph_format(heading_p, WD_ALIGN_PARAGRAPH.LEFT)
    heading_run = heading_p.add_run("1. Particulars")
    _set_font(heading_run, SECTION_HEADING_SIZE, bold=True)

    _add_labelled_paragraph(doc, "Case Study Title", particulars.case_study_title)
    _add_labelled_paragraph(doc, "Aim of Experiment", particulars.aim)
    _add_labelled_paragraph(doc, "Problem Statement", particulars.problem_statement)
    _add_labelled_paragraph(doc, "Author", particulars.author)
    _add_labelled_paragraph(doc, "Section", particulars.section)
    _add_labelled_paragraph(doc, "Roll Number", particulars.roll_number)
    _add_labelled_paragraph(
        doc, "Date of Compilation", particulars.date_of_compilation.strftime("%B %d, %Y")
    )


def _render_section(doc, section: Section):
    heading_p = doc.add_paragraph()
    _apply_paragraph_format(heading_p, WD_ALIGN_PARAGRAPH.LEFT)
    heading_run = heading_p.add_run(section.heading)
    _set_font(heading_run, SECTION_HEADING_SIZE, bold=True)

    for block in section.content:
        _render_block(doc, block)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def render_docx(report: ExperimentReport, output_path: str) -> str:
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # --- page setup ---
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM

    _add_page_number_footer(section)

    # --- document-level default style, as a safety net for any paragraph
    # that bypasses our explicit run-level font calls (e.g. list bullet
    # markers themselves) ---
    normal_style = doc.styles["Normal"]
    normal_style.font.name = FONT_NAME
    normal_style.font.size = NORMAL_SIZE

    # --- title ---
    title_p = doc.add_paragraph()
    _apply_paragraph_format(title_p, WD_ALIGN_PARAGRAPH.LEFT)
    title_run = title_p.add_run(report.title.upper())
    _set_font(title_run, TITLE_SIZE, bold=True)

    # --- particulars (section 1, always prose) ---
    _render_particulars(doc, report.particulars)

    # --- body sections (2, 3, 4, ...) ---
    for section_obj in report.sections:
        _render_section(doc, section_obj)

    doc.save(output_path)
    return output_path


def render_pdf(report: ExperimentReport, output_dir: str) -> str:
    """Renders DOCX then converts to PDF via headless LibreOffice, and
    names the final file per the guideline's naming convention
    (SubmissionMeta.filename())."""
    import subprocess
    import os
    import tempfile

    os.makedirs(output_dir, exist_ok=True)

    with tempfile.TemporaryDirectory(dir=output_dir) as tmp_dir:
        docx_path = os.path.join(tmp_dir, "render.docx")
        render_docx(report, docx_path)

        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                tmp_dir,
                docx_path,
            ],
            check=True,
            capture_output=True,
        )

        generated_pdf = os.path.join(tmp_dir, "render.pdf")
        final_name = report.submission_meta.filename(ext="pdf")
        final_path = os.path.join(output_dir, final_name)
        os.replace(generated_pdf, final_path)

    return final_path
